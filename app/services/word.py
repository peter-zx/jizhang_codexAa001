import logging

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.models import Person, StoredFile, User
from app.services.excel import DEFAULT_EXPORT_FIELDS, FIELD_DEFINITIONS
from app.services.storage import export_path

logger = logging.getLogger(__name__)


def _value(person: Person, field: str) -> str:
    if field == "owner":
        return person.owner.display_name
    if field == "monthly_confirmed":
        return "已核准" if person.monthly_confirmed else "待核准"
    if field == "confirmed_at":
        return person.confirmed_at.strftime("%Y-%m-%d %H:%M") if person.confirmed_at else ""
    value = getattr(person, field, "")
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def _set_document_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(9)
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)


def _add_field_line(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    label_run = paragraph.add_run(f"{label}：")
    label_run.bold = True
    label_run.font.name = "Microsoft YaHei"
    label_run.font.size = Pt(8.5)
    value_run = paragraph.add_run(value or " ")
    value_run.font.name = "Microsoft YaHei"
    value_run.font.size = Pt(8.5)


def _add_person_page(document: Document, person: Person, field_defs: list[tuple[str, str]], first: bool) -> None:
    if not first:
        document.add_section(WD_SECTION.NEW_PAGE)
    section = document.sections[-1]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("个人信息")
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(16)
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    sub = subtitle.add_run(person.name)
    sub.font.name = "Microsoft YaHei"
    sub.font.size = Pt(10)
    sub.font.color.rgb = RGBColor(102, 112, 133)

    for field, label in field_defs:
        _add_field_line(document, label, _value(person, field))


def export_people_to_docx(db, people: list[Person], current_user: User, period: str, fields: list[str] | None = None) -> StoredFile:
    try:
        allowed = dict(FIELD_DEFINITIONS)
        selected = [field for field in (fields or DEFAULT_EXPORT_FIELDS) if field in allowed and field != "service_fee"]
        if not selected:
            selected = DEFAULT_EXPORT_FIELDS
        field_defs = [(field, allowed[field]) for field in selected]

        document = Document()
        _set_document_style(document)
        if people:
            for index, person in enumerate(people):
                _add_person_page(document, person, field_defs, first=index == 0)
        else:
            document.add_heading("个人信息", level=0)
            document.add_paragraph("暂无人员资料")

        stored_name, path = export_path("个人信息", ".docx")
        document.save(path)
        record = StoredFile(owner_id=current_user.id, file_type="export", original_name=None, stored_name=stored_name, path=str(path))
        db.add(record)
        db.commit()
        return record
    except Exception:
        db.rollback()
        logger.exception("导出 Word 失败")
        raise
